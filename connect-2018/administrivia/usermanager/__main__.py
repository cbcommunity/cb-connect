#!/usr/bin/env python

#
# This script will take a CSV file containing attendees for the Cb Connect 2018 Developer Day conference
# and add the appropriate users and user groups into a Cb Response and Cb Protection server automatically
#

import sys
from cbapi.example_helpers import get_cb_response_object, get_cb_protection_object, build_cli_parser
import cbapi.response.models as cbr_models
import cbapi.protection.models as cbp_models
from cbapi.errors import ServerError, ApiError
import csv
import random
import string
import hashlib
import docx


def create_cbp_hash(password):
    candidate_letters = string.ascii_letters + string.digits
    salt = ''.join(random.choice(candidate_letters) for _ in range(10))

    passwordbytes = password + salt
    hashed_password = hashlib.md5(passwordbytes.encode('ascii')).hexdigest()

    return salt, hashed_password


def get_user_list(users):
    # For each user, create a username based on the password, strip off everything past the '@' sign
    response_users = {}
    for user in users:
        atsign = user[0].find('@')
        if atsign != -1:
            username = user[0][:atsign]
        else:
            username = user[0]
        email = user[0]
        password = user[1]

        if username in response_users:
            print("Warning: username {0} duplicated with email {1}".format(username, email))
        else:
            response_users[username] = (email, password)

    return response_users


def add_to_response(cbr, teamname, users):
    """Add the users to the Cb Response server"""

    # For each user, create a username based on the password, strip off everything past the '@' sign
    response_users = get_user_list(users)

    # Get list of users from Cb Protection, so we don't create a user who already exists
    existing_users = set([u.username for u in cbr.select(cbr_models.User)])
    new_users = set(response_users.keys())

    already_created_users = existing_users.intersection(new_users)

    for username in already_created_users:
        print("Warning: username {0} already exists, not creating".format(username))
        del response_users[username]

    # Get team
    team = cbr.select(cbr_models.Team).where("name:{0}".format(teamname)).first()
    if not team:
        print("Fatal: cannot find team named {0}".format(teamname))
        return False

    # Create the new users
    for username in response_users.keys():
        email, password = response_users[username]
        new_user = cbr.create(cbr_models.User)
        new_user.username = username
        new_user.email = email
        new_user.add_team(team)
        new_user.global_admin = True              # for now, we will make everyone global admin
        new_user.first_name = username
        new_user.password = password
        new_user.last_name = "user"

        try:
            new_user.save()
            print("Added user {0} to Cb Response".format(username))
        except ServerError as se:
            print("Warning: could not add user {0}: {1}".format(username, se))

    return True


def add_to_protection(cbp, groupname, users):
    """Add the users to the Cb Protection server"""

    # For each user, create a username based on the password, strip off everything past the '@' sign
    protection_users = get_user_list(users)

    # Get list of users from Cb Response, so we don't create a user who already exists
    existing_users = set([u.name for u in cbp.select(cbp_models.User)])
    new_users = set(protection_users.keys())

    already_created_users = existing_users.intersection(new_users)

    for username in already_created_users:
        print("Warning: username {0} already exists, not creating".format(username))
        del protection_users[username]

    # Get team
    team = cbp.select(cbp_models.UserGroup).where("name:{0}".format(groupname)).first()
    if not team:
        print("Fatal: cannot find team named {0}".format(groupname))
        return False

    # Create the new users
    for username in protection_users.keys():
        email, password = protection_users[username]
        new_user = cbp.create(cbp_models.User)
        new_user.name = username
        new_user.eMailAddress = email
        new_user.userGroupIds = team.id
        new_user.enabled = True
        new_user.unified = False
        new_user.passwordSalt, new_user.passwordHash = create_cbp_hash(password)

        try:
            new_user.save()
            print("Added user {0} to Cb Protection".format(username))
        except ServerError as se:
            print("Warning: could not add user {0}: {1}".format(username, se))

    return True


def read_from_csv(fn):
    """Read in the CSV file and return a list of (email, password) tuples"""
    users = []
    csvfile = csv.reader(open(fn, "r"))
    for row in csvfile:
        if row[0] != "":
            users.append(row)
    return users


def generate_user_credential_doc(users, args):
    d = docx.Document()

    userlist = get_user_list(users)

    for username in userlist.keys():
        email, password = userlist[username]

        d.add_heading("Login information", 0)

        d.add_paragraph("Welcome to Cb Connect Developer Day 2018! This page includes the login details so you can "
                        "access a Cb Response, Cb Defense, Cb Protection, and Cb ThreatHunter console.")
        d.add_paragraph("While your credentials are unique and not to be shared with anyone else in the conference, "
                        "please be aware that these are shared resources and please do not intentionally perform "
                        "any queries to degrade or deny service to others, or to perform administrative tasks that "
                        "will affect sensors attached to these instances or other user accounts.")

        d.add_heading("Cb Response", 1)

        d.add_paragraph("Login URL to Cb Response: {0}".format(args.response))
        d.add_paragraph("Username: {0}".format(username))
        d.add_paragraph("Password: {0}".format(password))

        d.add_heading("Cb Protection", 1)

        d.add_paragraph("Login URL to Cb Protection: {0}".format(args.protection))
        d.add_paragraph("Username: {0}".format(username))
        d.add_paragraph("Password: {0}".format(password))

        d.add_heading("Cb Defense & ThreatHunter (PSC)", 1)

        d.add_paragraph("Login URL to PSC: {0}".format(args.defense))
        d.add_paragraph("Username: {0}".format(email))
        last_para = d.add_paragraph("Password: {0}".format(password))

        last_para.runs[0].add_break(docx.text.run.WD_BREAK.PAGE)

    return d


def main():
    parser = build_cli_parser("Manage users on Cb Connect 2018 Developer Day servers")
    commands = parser.add_subparsers(help="User commands", dest="command_name")

    bulk_command = commands.add_parser("bulkadd", help="Bulk add users from CSV file")
    bulk_command.add_argument("-f", "--csvfile", help="input CSV file name", required=True)
    bulk_command.add_argument("-g", "--group", help="Group to add these users to", default="devday")

    print_command = commands.add_parser("print", help="Print user credentials to docx file")
    print_command.add_argument("-p", "--protection", help="Cb Protection base login URL",
                               default="https://cbprotection.devday2018.com")
    print_command.add_argument("-d", "--defense", help="Cb Defense base login URL",
                               default="https://defense-eap01.conferdeploy.net")
    print_command.add_argument("-r", "--response", help="Cb Response base login URL",
                               default="https://cbresponse.devday2018.com")
    print_command.add_argument("-o", "--output", help="Output file for credentials docx file",
                               default="credentials.docx")
    print_command.add_argument("-f", "--csvfile", help="input CSV file name", required=True)

    args = parser.parse_args()

    if args.command_name == "bulkadd":
        # Get the Cb Response and Cb Protection API objects, respectively
        cbr = get_cb_response_object(args)
        cbp = get_cb_protection_object(args)

        # Read the list of emails and passwords from our CSV file
        users = read_from_csv(args.csvfile)

        # Create the Response users
        add_to_response(cbr, args.group, users)
        add_to_protection(cbp, args.group, users)
    elif args.command_name == "print":
        # Read the list of emails and passwords from our CSV file
        users = read_from_csv(args.csvfile)

        doc = generate_user_credential_doc(users, args)
        doc.save(args.output)
    else:
        print("Invalid command: {0}".format(args.command_name))


if __name__ == '__main__':
    sys.exit(main())